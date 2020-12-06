import openpyxl
from docx import Document


def rating_local(p, p_max, r, r_max):
    a_1 = 1 / 3
    a_2 = 1 - a_1
    return a_1 * p / p_max + a_2 * r / r_max


def find_best_suppliers(prod, data):
    wb = openpyxl.load_workbook(data)

    ws = wb["product"]
    prod_id = ''
    for row in ws.rows:
        if row[1].value == prod:
            prod_id = row[0].value
            # print(prod_id)
            break

    ws = wb["price"]
    suppliers = {}
    p_max = 0
    for row in ws.rows:
        if row[1].value == prod_id:
            supplier_id, prod_id, price, term = [c.value for c in row]
            price = float(price.replace(',', '.'))
            if price > p_max:
                p_max = price
            suppliers[supplier_id] = [price, prod_id, prod]
    # print(suppliers, p_max)

    ws = wb["suppliers"]
    r_max = 0
    for row in ws.rows:
        if row[0].value in suppliers.keys():
            supplier_id, name, rating, adress = [c.value for c in row]
            # rating = float(rating.replace(',', '.'))
            if rating > r_max:
                r_max = rating
            suppliers[supplier_id] += [name, rating, adress]
    # print(suppliers)

    to_sort = []
    for supplier_id in suppliers.keys():
        p = suppliers[supplier_id][0]
        r = suppliers[supplier_id][4]
        to_sort.append((supplier_id, rating_local(p, p_max, r, r_max)))
    from_sort = sorted(to_sort, key=lambda item: item[1], reverse=True)
    from_sort = from_sort[:3]
    # print(from_sort)

    result = {}
    for sup in from_sort:
        supplier_id = sup[0]
        result[supplier_id] = suppliers[supplier_id]

    return result


def find_best_suppliers_list(data, prod_list):
    output_list = []
    for prod in prod_list:
        output_list.append(find_best_suppliers(prod, data))

    return output_list


def write_letters(data, prod_list):
    write_to = {}
    for prod in find_best_suppliers_list(data, prod_list):
        for sup_id, info in prod.items():
            if sup_id not in write_to.keys():
                write_to[sup_id] = [info, ]
            else:
                write_to[sup_id].append(info)
    # print(write_to)

    for prod_id, prods in write_to.items():
        address = prods[0][5]
        letter_name = "to_" + address + ".docx"

        doc = Document("template.docx")
        for prod in prods:
            prod_name = prod[2]
            p = doc.add_paragraph(style='List')
            r = p.add_run()
            r.add_text(prod_name)

        doc.save("letters/" + letter_name)


if __name__ == '__main__':
    prod_list = ['Олівець', "Ручка кулькова"]
    data = "data.xlsx"
    write_letters(data, prod_list)

